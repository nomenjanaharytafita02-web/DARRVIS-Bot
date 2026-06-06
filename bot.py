import os
import logging
import requests
import json

from fpdf import FPDF
from docx import Document as DocxDocument
from PyPDF2 import PdfReader

from telegram import Update
from telegram.constants import ChatAction
from telegram.ext import Application, CommandHandler, MessageHandler, filters, ContextTypes

# Configuration
TELEGRAM_TOKEN = os.getenv("TELEGRAM_TOKEN")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
BOT_NAME = "Darrvis"

GEMINI_API_URL = "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash:generateContent"
POLLINATIONS_API_URL = "https://image.pollinations.ai/prompt/"

WEBHOOK_URL = os.getenv("WEBHOOK_URL")
PORT = int(os.getenv("PORT", 8000))

# Set up logging
logging.basicConfig(format="%(asctime)s - %(name)s - %(levelname)s - %(message)s", level=logging.INFO)
logger = logging.getLogger(__name__)

# In-memory storage for conversation context and message count
user_conversations = {}
user_message_counts = {}
MAX_FREE_MESSAGES = 10

# --- Utility Functions ---

def get_gemini_response(prompt, chat_history=None):
    headers = {"Content-Type": "application/json"}
    payload = {"contents": []}

    if chat_history:
        for message in chat_history:
            payload["contents"].append({"role": message["role"], "parts": [{"text": message["text"]}]})

    payload["contents"].append({"role": "user", "parts": [{"text": prompt}]})

    params = {"key": GEMINI_API_KEY}
    response = requests.post(GEMINI_API_URL, headers=headers, json=payload, params=params)
    response.raise_for_status()
    return response.json()["candidates"][0]["content"]["parts"][0]["text"]

# --- Telegram Bot Handlers ---

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.effective_chat.send_chat_action(ChatAction.TYPING)
    user_id = update.effective_user.id
    user_conversations[user_id] = []
    user_message_counts[user_id] = 0
    welcome_message = (
        f"Salut ! Je suis {BOT_NAME}, votre assistant IA polyvalent.\n\n"
        "Je peux repondre a vos questions, generer des images, creer des documents PDF et Word, "
        "lire des fichiers, ecrire du code, traduire des langues, et bien plus encore !\n\n"
        "Tapez /aide pour voir la liste de mes fonctionnalites.\n"
        "N'hesitez pas a commencer a discuter avec moi !"
    )
    await update.message.reply_text(welcome_message)

async def aide(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.effective_chat.send_chat_action(ChatAction.TYPING)
    help_message = (
        "Voici ce que je peux faire :\n\n"
        "1. Chat IA (Gemini) : Reponds a toutes vos questions et garde le contexte de la conversation.\n"
        "2. Generation d'images : Utilisez /image [description] pour generer une image.\n"
        "3. Creation de PDF : Utilisez /pdf [sujet] pour creer un document PDF.\n"
        "4. Creation de Word : Utilisez /word [sujet] pour creer un document Word.\n"
        "5. Lecture de fichiers : Envoyez-moi un fichier PDF ou Word, et je pourrai en extraire le texte.\n"
        "6. Code : Je peux ecrire, corriger ou expliquer du code.\n"
        "7. Langues : Je peux traduire du texte dans toutes les langues.\n"
        "8. Monetisation : Vous avez 10 messages gratuits par jour.\n\n"
        "Commandes disponibles :\n"
        "/start - Message de bienvenue\n"
        "/aide - Cette liste de fonctionnalites\n"
        "/image [description] - Generer une image\n"
        "/pdf [sujet] - Creer un PDF\n"
        "/word [sujet] - Creer un Word\n"
        "/reset - Reinitialiser la conversation\n"
    )
    await update.message.reply_text(help_message)

async def reset(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.effective_chat.send_chat_action(ChatAction.TYPING)
    user_id = update.effective_user.id
    user_conversations[user_id] = []
    user_message_counts[user_id] = 0
    await update.message.reply_text("Votre conversation et votre compteur de messages ont ete reinitialises.")

async def handle_message(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.effective_chat.send_chat_action(ChatAction.TYPING)
    user_id = update.effective_user.id
    text = update.message.text

    if user_id not in user_message_counts:
        user_message_counts[user_id] = 0
    if user_id not in user_conversations:
        user_conversations[user_id] = []

    if user_message_counts[user_id] >= MAX_FREE_MESSAGES:
        await update.message.reply_text(
            "Vous avez atteint la limite de 10 messages gratuits par jour. "
            "Veuillez recharger votre compte pour continuer."
        )
        return

    user_message_counts[user_id] += 1
    user_conversations[user_id].append({"role": "user", "text": text})

    try:
        gemini_response = get_gemini_response(text, user_conversations[user_id])
        user_conversations[user_id].append({"role": "model", "text": gemini_response})
        await update.message.reply_text(gemini_response)
    except Exception as e:
        logger.error(f"Error getting Gemini response: {e}")
        await update.message.reply_text("Desole, une erreur est survenue lors de la communication avec l'IA.")

async def generate_image_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.effective_chat.send_chat_action(ChatAction.TYPING)
    user_id = update.effective_user.id

    if user_id not in user_message_counts:
        user_message_counts[user_id] = 0

    if user_message_counts[user_id] >= MAX_FREE_MESSAGES:
        await update.message.reply_text(
            "Vous avez atteint la limite de 10 messages gratuits par jour."
        )
        return

    user_message_counts[user_id] += 1

    if not context.args:
        await update.message.reply_text("Veuillez fournir une description. Exemple: /image un chat volant")
        return

    description = " ".join(context.args)
    image_url = f"{POLLINATIONS_API_URL}{description}"
    await update.message.reply_photo(photo=image_url, caption=f"Voici votre image pour: {description}")

async def create_pdf_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.effective_chat.send_chat_action(ChatAction.TYPING)
    user_id = update.effective_user.id

    if user_id not in user_message_counts:
        user_message_counts[user_id] = 0

    if user_message_counts[user_id] >= MAX_FREE_MESSAGES:
        await update.message.reply_text("Vous avez atteint la limite de 10 messages gratuits par jour.")
        return

    user_message_counts[user_id] += 1

    if not context.args:
        await update.message.reply_text("Veuillez fournir un sujet. Exemple: /pdf un rapport sur l'IA")
        return

    subject = " ".join(context.args)
    file_name = f"rapport_{user_id}.pdf"

    try:
        gemini_prompt = f"Genere un contenu detaille pour un document PDF sur le sujet suivant : {subject}. Le contenu doit etre bien structure avec des titres et des paragraphes. Utilise le francais."
        content = get_gemini_response(gemini_prompt)

        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.multi_cell(0, 10, content.encode("latin-1", "replace").decode("latin-1"))
        pdf.output(file_name)

        with open(file_name, "rb") as f:
            await update.message.reply_document(document=f, caption=f"Voici votre document PDF sur: {subject}")
        os.remove(file_name)
    except Exception as e:
        logger.error(f"Error creating PDF: {e}")
        await update.message.reply_text("Desole, une erreur est survenue lors de la creation du PDF.")

async def create_word_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.effective_chat.send_chat_action(ChatAction.TYPING)
    user_id = update.effective_user.id

    if user_id not in user_message_counts:
        user_message_counts[user_id] = 0

    if user_message_counts[user_id] >= MAX_FREE_MESSAGES:
        await update.message.reply_text("Vous avez atteint la limite de 10 messages gratuits par jour.")
        return

    user_message_counts[user_id] += 1

    if not context.args:
        await update.message.reply_text("Veuillez fournir un sujet. Exemple: /word une lettre de motivation")
        return

    subject = " ".join(context.args)
    file_name = f"document_{user_id}.docx"

    try:
        gemini_prompt = f"Genere un contenu detaille pour un document Word sur le sujet suivant : {subject}. Le contenu doit etre bien structure avec des titres et des paragraphes. Utilise le francais."
        content = get_gemini_response(gemini_prompt)

        document = DocxDocument()
        document.add_heading(subject, level=1)
        for paragraph in content.split("\n\n"):
            document.add_paragraph(paragraph)
        document.save(file_name)

        with open(file_name, "rb") as f:
            await update.message.reply_document(document=f, caption=f"Voici votre document Word sur: {subject}")
        os.remove(file_name)
    except Exception as e:
        logger.error(f"Error creating Word document: {e}")
        await update.message.reply_text("Desole, une erreur est survenue lors de la creation du document Word.")

async def handle_document(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.effective_chat.send_chat_action(ChatAction.TYPING)
    user_id = update.effective_user.id

    if user_id not in user_message_counts:
        user_message_counts[user_id] = 0

    if user_message_counts[user_id] >= MAX_FREE_MESSAGES:
        await update.message.reply_text("Vous avez atteint la limite de 10 messages gratuits par jour.")
        return

    user_message_counts[user_id] += 1

    doc = update.message.document
    file_name = doc.file_name

    if file_name.endswith(".pdf"):
        file_obj = await doc.get_file()
        downloaded_file = await file_obj.download_to_drive()

        try:
            reader = PdfReader(str(downloaded_file))
            text_content = ""
            for page in reader.pages:
                text_content += page.extract_text() + "\n"

            if user_id not in user_conversations:
                user_conversations[user_id] = []
            user_conversations[user_id].append({"role": "user", "text": f"Contenu du PDF {file_name}:\n{text_content[:2000]}"})
            await update.message.reply_text(f"J'ai lu le document PDF \"{file_name}\". Vous pouvez maintenant me poser des questions a ce sujet.")
        except Exception as e:
            logger.error(f"Error reading PDF: {e}")
            await update.message.reply_text("Desole, une erreur est survenue lors de la lecture du PDF.")
        finally:
            os.remove(str(downloaded_file))

    elif file_name.endswith(".docx"):
        file_obj = await doc.get_file()
        downloaded_file = await file_obj.download_to_drive()

        try:
            docx_doc = DocxDocument(str(downloaded_file))
            text_content = "\n".join([p.text for p in docx_doc.paragraphs])

            if user_id not in user_conversations:
                user_conversations[user_id] = []
            user_conversations[user_id].append({"role": "user", "text": f"Contenu du Word {file_name}:\n{text_content[:2000]}"})
            await update.message.reply_text(f"J'ai lu le document Word \"{file_name}\". Vous pouvez maintenant me poser des questions a ce sujet.")
        except Exception as e:
            logger.error(f"Error reading Word document: {e}")
            await update.message.reply_text("Desole, une erreur est survenue lors de la lecture du document Word.")
        finally:
            os.remove(str(downloaded_file))
    else:
        await update.message.reply_text("Desole, je ne peux lire que les fichiers PDF et Word pour le moment.")

# --- Main: Native webhook from python-telegram-bot ---

def main():
    """Start the bot using python-telegram-bot's native webhook server."""
    application = Application.builder().token(TELEGRAM_TOKEN).build()

    # Command Handlers
    application.add_handler(CommandHandler("start", start))
    application.add_handler(CommandHandler("aide", aide))
    application.add_handler(CommandHandler("reset", reset))
    application.add_handler(CommandHandler("image", generate_image_command))
    application.add_handler(CommandHandler("pdf", create_pdf_command))
    application.add_handler(CommandHandler("word", create_word_command))

    # Message Handlers
    application.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_message))
    application.add_handler(MessageHandler(filters.Document.ALL, handle_document))

    logger.info(f"Starting bot with webhook on port {PORT}...")

    if WEBHOOK_URL:
        # Use native webhook - no Flask, no gunicorn needed
        application.run_webhook(
            listen="0.0.0.0",
            port=PORT,
            url_path=TELEGRAM_TOKEN,
            webhook_url=f"{WEBHOOK_URL}/{TELEGRAM_TOKEN}",
        )
    else:
        # Fallback to polling for local development
        application.run_polling(drop_pending_updates=True)

if __name__ == "__main__":
    main()
